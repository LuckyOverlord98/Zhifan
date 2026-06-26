from __future__ import annotations

import argparse
import json
import re
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
DATA_PATH = ROOT / "data" / "jinqiao-products.json"
DOCX_PATH = ROOT / "金桥待补充产品介绍与适用场景.docx"

GENERIC_INTRO = "详情页整理执行标准"
GENERIC_APPLICATION = "按母材、强度等级"

MODEL_RENAMES = {
    "JQ.TH500-NQ-II": "JQ.TH500-NQ",
    "JQ.TH550-NQ-II": "JQ.TH550-NQ",
    "JQ.TH650EW-II": "JQ.TH650EW",
    "JQ·H08MnSiCuCrNi-Ⅱ": "JQ.H08MnSiCuCrNi",
}

CATEGORY_SLUGS = {
    "碳钢焊条": "carbon-steel-electrodes",
    "实芯气保及氩弧焊丝": "solid-wires",
    "药芯气保焊丝": "flux-cored-wires",
    "不锈钢焊材": "stainless-materials",
    "碳钢埋弧焊丝焊剂": "submerged-arc",
    "铝焊丝": "aluminum-wires",
    "特种焊材": "special-materials",
    "设备配件与工具": "equipment-accessories",
}


def normalize_model(value: str) -> str:
    value = (value or "").strip().upper()
    value = value.replace("·", ".").replace("•", ".").replace("．", ".").replace("。", ".")
    value = value.replace("－", "-").replace("–", "-").replace("—", "-").replace("Ⅱ", "II")
    value = value.replace("（", "(").replace("）", ")")
    value = re.sub(r"^JQ[.\-]?", "", value)
    value = re.sub(r"[^A-Z0-9]+", "", value)
    return value


def slugify_model(model: str) -> str:
    slug = model.strip().lower()
    slug = slug.replace("·", "-").replace(".", "-").replace("（", "-").replace("）", "")
    slug = slug.replace("(", "-").replace(")", "")
    slug = slug.replace("Ⅱ", "ii")
    slug = re.sub(r"[^a-z0-9]+", "-", slug)
    return slug.strip("-")


def clean_cell(value: str) -> str:
    value = (value or "").replace("\u3000", " ")
    value = re.sub(r"[ \t]*\n[ \t]*", "", value)
    value = re.sub(r"\s+", " ", value)
    value = value.replace(" 。", "。").replace(" ，", "，").strip()
    return value.rstrip(" /")


def clean_intro(value: str) -> str:
    value = clean_cell(value)
    value = re.sub(r"^(特性|说明)[:：]", "", value).strip()
    return value


def split_applications(value: str) -> list[str]:
    value = clean_cell(value)
    value = re.sub(r"^(用途|适用场景)[:：]", "", value).strip()
    pieces = []
    for part in re.split(r"[。；;]", value):
        part = part.strip(" ，,。")
        if not part:
            continue
        part = re.sub(r"^(用途|特性)[:：]", "", part).strip()
        if not part:
            continue
        pieces.append(part)

    result = []
    seen = set()
    for piece in pieces:
        if not piece.startswith(("适用于", "常见应用", "主要应用", "可用于", "用于")):
            piece = "适用于" + piece
        key = piece.replace(" ", "")
        if key in seen:
            continue
        seen.add(key)
        result.append(piece)
    return result[:6]


def split_standards(value: str) -> list[str]:
    value = clean_cell(value)
    result = []
    seen = set()
    for part in value.split("/"):
        part = part.strip()
        if not part or part == "--":
            continue
        key = part.replace(" ", "")
        if key in seen:
            continue
        seen.add(key)
        result.append(part)
    return result


def load_docx_rows() -> list[dict]:
    doc = Document(DOCX_PATH)
    if not doc.tables:
        raise RuntimeError(f"No tables found in {DOCX_PATH}")
    rows = []
    for row in doc.tables[0].rows[1:]:
        cells = [cell.text.strip() for cell in row.cells]
        if len(cells) < 7:
            continue
        model = cells[1].strip()
        category = clean_cell(cells[2])
        standards = split_standards(cells[3])
        intro = clean_intro(cells[5])
        applications = split_applications(cells[6])
        if not model or not intro or not applications:
            continue
        rows.append({
            "displayModel": model,
            "matchModel": MODEL_RENAMES.get(model, model),
            "categoryName": category,
            "categorySlug": CATEGORY_SLUGS.get(category, ""),
            "standards": standards,
            "intro": intro,
            "applications": applications,
        })
    return rows


def product_lookup(products: list[dict]) -> dict[str, dict]:
    lookup = {}
    for product in products:
        if product.get("manufacturer") != "金桥":
            continue
        lookup.setdefault(normalize_model(product.get("model", "")), product)
    return lookup


def apply_rows(products: list[dict], rows: list[dict]) -> dict:
    lookup = product_lookup(products)
    updated = []
    created = []
    for row in rows:
        product = lookup.get(normalize_model(row["matchModel"])) or lookup.get(normalize_model(row["displayModel"]))
        display_model = row["displayModel"].replace("·", ".").replace("（", "(").replace("）", ")").replace("Ⅱ", "II")
        if display_model.startswith("H08"):
            display_model = "JQ." + display_model

        if not product:
            product = {
                "slug": slugify_model(display_model),
                "manufacturer": "金桥",
                "categorySlug": row["categorySlug"] or "solid-wires",
                "categoryName": row["categoryName"] or "实芯气保及氩弧焊丝",
                "model": display_model,
                "name": f"金桥 {display_model} {row['categoryName']}".strip(),
                "standard": " / ".join(row["standards"]),
                "standards": row["standards"],
                "summary": f"金桥 {display_model}，适用于{row['categoryName']}相关工况。",
                "introduction": row["intro"],
                "applications": row["applications"],
                "composition": [],
                "depositedMetal": [],
                "dimensions": [],
                "notes": "",
                "source": "金桥待补充产品介绍与适用场景.docx",
            }
            products.append(product)
            lookup[normalize_model(display_model)] = product
            created.append(display_model)
            continue

        old_model = product.get("model", "")
        product["model"] = display_model
        product["slug"] = slugify_model(display_model)
        product["name"] = f"金桥 {display_model} {product.get('categoryName', '').strip()}".strip()
        if product.get("summary"):
            product["summary"] = f"金桥 {display_model}，适用于{product.get('categoryName', '')}相关工况。"
        product["introduction"] = row["intro"]
        product["applications"] = row["applications"]
        updated.append({"from": old_model, "to": display_model})

    return {"updated": updated, "created": created}


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--write", action="store_true")
    args = parser.parse_args()

    products = json.loads(DATA_PATH.read_text(encoding="utf-8-sig"))
    rows = load_docx_rows()
    result = apply_rows(products, rows)

    generic_left = [
        product.get("model")
        for product in products
        if product.get("manufacturer") == "金桥"
        and (
            GENERIC_INTRO in product.get("introduction", "")
            or any(GENERIC_APPLICATION in item for item in product.get("applications", []))
        )
    ]

    print(json.dumps({
        "docxRows": len(rows),
        "updated": result["updated"],
        "created": result["created"],
        "remainingGenericJinqiaoContent": generic_left,
    }, ensure_ascii=False, indent=2))

    if args.write:
        DATA_PATH.write_text(json.dumps(products, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")


if __name__ == "__main__":
    main()
