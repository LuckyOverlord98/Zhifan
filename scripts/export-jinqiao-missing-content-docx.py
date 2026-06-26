from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DATA_PATH = ROOT / "data" / "jinqiao-products.json"
OUT_PATH = ROOT / "金桥待补充产品介绍与适用场景.docx"

GENERIC_INTRO = "详情页整理执行标准"
GENERIC_APPLICATION = "按母材、强度等级"


def missing_products():
    products = json.loads(DATA_PATH.read_text(encoding="utf-8-sig"))
    rows = []
    for product in products:
        if product.get("manufacturer") != "金桥":
            continue
        generic_intro = GENERIC_INTRO in product.get("introduction", "")
        generic_apps = any(GENERIC_APPLICATION in item for item in product.get("applications", []))
        if generic_intro or generic_apps:
            rows.append(product)
    return rows


def set_cell_text(cell, text, bold=False, color=None, size=8.5):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text or "")
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def apply_table_widths(table, widths):
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = Inches(width)


def main():
    rows = missing_products()

    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.6)
    section.right_margin = Inches(0.6)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(10)

    title = doc.add_paragraph()
    title_run = title.add_run("金桥产品介绍与适用场景待补充清单")
    title_run.font.size = Pt(18)
    title_run.bold = True
    title_run.font.color.rgb = RGBColor(31, 77, 120)

    subtitle = doc.add_paragraph()
    subtitle.add_run(
        f"共 {len(rows)} 个型号未在已提取手册文本中稳定匹配到“用途/特性”完整内容。"
        "请补充产品介绍、适用场景和必要备注，后续可按本表回填网站数据库。"
    )

    table = doc.add_table(rows=1, cols=7)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    headers = ["序号", "型号", "产品分类", "执行标准", "当前状态", "产品介绍待填", "适用场景待填 / 备注"]
    widths = [0.45, 1.35, 1.35, 2.25, 1.2, 2.35, 2.45]
    for cell, header in zip(table.rows[0].cells, headers):
        set_cell_text(cell, header, bold=True, color="1F4D78", size=8.5)

    for index, product in enumerate(rows, 1):
        row = table.add_row().cells
        status = []
        if GENERIC_INTRO in product.get("introduction", ""):
            status.append("缺产品介绍")
        if any(GENERIC_APPLICATION in item for item in product.get("applications", [])):
            status.append("缺适用场景")
        values = [
            str(index),
            product.get("model", ""),
            product.get("categoryName", ""),
            product.get("standard", ""),
            "；".join(status),
            "",
            "",
        ]
        for cell, value in zip(row, values):
            set_cell_text(cell, value, size=8)

    apply_table_widths(table, widths)

    note = doc.add_paragraph()
    note.add_run("填写建议：").bold = True
    note.add_run(" 产品介绍优先写材料体系、药皮/焊丝类型、工艺性能；适用场景优先写强度级别、母材牌号、行业或典型项目。")

    doc.save(OUT_PATH)
    print(OUT_PATH)


if __name__ == "__main__":
    main()
