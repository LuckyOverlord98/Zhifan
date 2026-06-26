from pathlib import Path
from docx import Document
import html
import json
import re


ROOT = Path(__file__).resolve().parents[1]
SOURCE_DIR = Path("C:/Users/Ning Sun/Documents/Codex/2026-06-25/new-chat/outputs")
GENERATED_AT = "2026-06-26"
SEO_KEYWORDS = [
    "宁波焊材批发",
    "浙江焊材供应商",
    "江浙沪焊材配送",
    "焊材选型",
    "焊接操作",
    "焊材质量证明书",
    "焊材现货供应",
    "金桥焊材",
    "大西洋焊材",
    "上海东风焊材",
    "天泰焊材",
    "船用焊材",
    "钢结构焊材",
    "压力容器焊材",
]


def find_source_docx():
    matches = list(SOURCE_DIR.glob("*80*.docx"))
    if not matches:
        raise SystemExit("No 80 article DOCX found")
    return matches[0]


def slugify(text, number):
    mapping = {
        "J422": "j422",
        "J507": "j507",
        "J506": "j506",
        "E4303": "e4303",
        "E5015": "e5015",
        "E5016": "e5016",
        "ER50-6": "er50-6",
        "ER70S-6": "er70s-6",
        "ER308L": "er308l",
        "ER316L": "er316l",
        "A102": "a102",
        "A107": "a107",
        "A132": "a132",
        "D256": "d256",
        "D507": "d507",
        "D707": "d707",
        "Q235": "q235",
        "Q345": "q345",
        "16Mn": "16mn",
        "304": "304",
        "316L": "316l",
    }
    keyword_map = [
        ("焊条", "electrode"),
        ("焊丝", "wire"),
        ("焊剂", "flux"),
        ("型号", "model"),
        ("牌号", "grade"),
        ("酸性", "acid"),
        ("碱性", "basic"),
        ("实芯", "solid"),
        ("药芯", "flux-cored"),
        ("气保", "gas-shielded"),
        ("氩弧", "tig"),
        ("埋弧", "saw"),
        ("执行标准", "standard"),
        ("低碳钢", "low-carbon-steel"),
        ("钢结构", "steel-structure"),
        ("船舶", "shipbuilding"),
        ("压力容器", "pressure-vessel"),
        ("管道", "pipeline"),
        ("锅炉", "boiler"),
        ("石化", "petrochemical"),
        ("电力", "power"),
        ("工程机械", "machinery"),
        ("矿山", "mining"),
        ("模具", "mold"),
        ("耐磨", "wear-resistant"),
        ("堆焊", "hardfacing"),
        ("气孔", "porosity"),
        ("裂纹", "crack"),
        ("夹渣", "slag"),
        ("飞溅", "spatter"),
        ("成型", "bead-shape"),
        ("粘条", "sticking"),
        ("咬边", "undercut"),
        ("强度", "strength"),
        ("变形", "distortion"),
        ("电弧", "arc"),
        ("烘干", "drying"),
        ("仓库", "warehouse"),
        ("生锈", "rust"),
        ("受潮", "moisture"),
        ("保温筒", "holding-oven"),
        ("合格证", "certificate"),
        ("质保书", "quality-certificate"),
        ("批次号", "batch"),
        ("采购", "purchase"),
        ("不锈钢", "stainless"),
        ("碳钢", "carbon-steel"),
    ]
    found = []
    lower = text.lower()
    for key, value in mapping.items():
        if key.lower() in lower:
            found.append(value)
    for key, value in keyword_map:
        if key in text:
            found.append(value)
    compact = "-".join(dict.fromkeys(found))
    return f"qa-{number}-{compact or 'welding'}"


def extract_articles(doc):
    paras = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    articles = []
    i = 0
    while i < len(paras):
        match = re.match(r"^(\d{2})[\.．、]\s*(.+)$", paras[i])
        if not match:
            i += 1
            continue
        number, title = match.group(1), match.group(2).strip()
        category = paras[i + 1].strip() if i + 1 < len(paras) else ""
        short_answer = ""
        long_answer = ""
        table_title = ""
        note = ""
        j = i + 2
        if j < len(paras) and "100" in paras[j]:
            j += 1
            short_answer = paras[j].strip() if j < len(paras) else ""
            j += 1
        if j < len(paras) and "500" in paras[j]:
            j += 1
            long_answer = paras[j].strip() if j < len(paras) else ""
            j += 1
        if j < len(paras) and not re.match(r"^(\d{2})[\.．、]", paras[j]):
            table_title = paras[j].strip()
            j += 1
        if j < len(paras) and not re.match(r"^(\d{2})[\.．、]", paras[j]):
            note = paras[j].strip()
            j += 1
        articles.append(
            {
                "number": number,
                "title": title,
                "category": category,
                "shortAnswer": short_answer,
                "longAnswer": long_answer,
                "tableTitle": table_title,
                "note": note,
            }
        )
        i = j

    if len(articles) != 80:
        raise SystemExit(f"Expected 80 articles, got {len(articles)}")
    if len(doc.tables) < len(articles):
        raise SystemExit(f"Expected >= 80 tables, got {len(doc.tables)}")

    for idx, article in enumerate(articles):
        rows = []
        for row in doc.tables[idx].rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                rows.append(cells)
        article["table"] = rows

    category_order = []
    category_slugs = {}
    slug_names = [
        "basics",
        "electrode-selection",
        "gas-shielded-wire",
        "stainless-materials",
        "hardfacing-special",
        "industry-applications",
        "welding-troubleshooting",
        "storage-drying",
    ]
    for article in articles:
        if article["category"] not in category_order:
            category_order.append(article["category"])
    for idx, category in enumerate(category_order):
        category_slugs[category] = slug_names[idx] if idx < len(slug_names) else f"category-{idx + 1}"

    for article in articles:
        article["categorySlug"] = category_slugs[article["category"]]
        article["slug"] = slugify(article["title"], article["number"])
        article["href"] = f"/articles/qa/{article['slug']}.html"
        article["summary"] = article["shortAnswer"][:92] + ("..." if len(article["shortAnswer"]) > 92 else "")
        article["generatedAt"] = GENERATED_AT

    categories = [
        {
            "slug": category_slugs[category],
            "title": category,
            "count": sum(1 for article in articles if article["category"] == category),
        }
        for category in category_order
    ]
    return articles, categories


def esc(value):
    return html.escape(str(value or ""), quote=True)


def paragraphize(text):
    text = str(text or "").strip()
    if not text:
        return ""
    parts = re.split(r"(?<=[。！？])\s*", text)
    chunks = []
    buffer = ""
    for part in parts:
        if not part:
            continue
        if len(buffer) + len(part) < 180:
            buffer += part
        else:
            if buffer:
                chunks.append(buffer)
            buffer = part
    if buffer:
        chunks.append(buffer)
    return "\n".join(f"<p>{esc(chunk)}</p>" for chunk in chunks)


def table_html(article):
    rows = article.get("table") or []
    if not rows:
        return ""
    header = rows[0]
    body = rows[1:]
    thead = "<thead><tr>" + "".join(f"<th>{esc(cell)}</th>" for cell in header) + "</tr></thead>"
    tbody = "<tbody>" + "".join(
        "<tr>" + "".join(f"<td>{esc(cell)}</td>" for cell in row) + "</tr>" for row in body
    ) + "</tbody>"
    title = esc(article.get("tableTitle") or "选型参考表")
    return f'<div class="article-table-wrap qa-table-wrap"><h2>{title}</h2><table class="article-table qa-table">{thead}{tbody}</table></div>'


def write_data_module(articles, categories):
    out = ROOT / "src" / "data" / "knowledgeQa.js"
    out.parent.mkdir(parents=True, exist_ok=True)
    out.write_text(
        "export const knowledgeQaCategories = "
        + json.dumps(categories, ensure_ascii=True, indent=2)
        + ";\n\nexport const knowledgeQaArticles = "
        + json.dumps(articles, ensure_ascii=True, indent=2)
        + ";\n",
        encoding="utf-8",
    )



def article_keywords(article):
    text = " ".join([
        article.get("title", ""),
        article.get("category", ""),
        article.get("shortAnswer", ""),
        article.get("longAnswer", ""),
    ])
    model_terms = re.findall(r"(?:[A-Za-z]{1,6}[A-Za-z0-9·.\-]*\d+[A-Za-z0-9·.\-]*|Q\d{3,4}|16Mn)", text)
    terms = [article.get("title", ""), article.get("category", ""), *model_terms, *SEO_KEYWORDS]
    result = []
    seen = set()
    for term in terms:
        term = re.sub(r"\s+", " ", str(term or "")).strip()
        key = term.lower().replace(" ", "")
        if not term or key in seen:
            continue
        seen.add(key)
        result.append(term)
    return ",".join(result[:32])


def article_picture_html():
    return '<picture><source type="image/webp" srcset="../../assets/optimized/sections__knowledge-operation-480.webp 480w, ../../assets/optimized/sections__knowledge-operation-768.webp 768w, ../../assets/optimized/sections__knowledge-operation-1280.webp 1280w" sizes="(max-width: 760px) 100vw, 48vw" /><img src="../../assets/sections/knowledge-operation.png" alt="焊接操作与焊材选型" loading="lazy" decoding="async" /></picture>'
def write_pages(articles):
    out_dir = ROOT / "public" / "articles" / "qa"
    out_dir.mkdir(parents=True, exist_ok=True)
    for old in out_dir.glob("*.html"):
        old.unlink()

    for idx, article in enumerate(articles):
        prev_link = articles[idx - 1]["href"] if idx > 0 else "../../index.html#knowledge"
        prev_text = "上一篇" if idx > 0 else "返回焊接操作"
        next_link = articles[idx + 1]["href"] if idx + 1 < len(articles) else "../../index.html#knowledge"
        next_text = "下一篇" if idx + 1 < len(articles) else "返回焊接操作"
        note = (
            f'<div class="qa-note"><strong>资料提示</strong><p>{esc(article["note"])}</p></div>'
            if article.get("note")
            else ""
        )
        page = f'''<!doctype html>
<html lang="zh-CN">
  <head>
    <meta charset="UTF-8" />
    <meta name="viewport" content="width=device-width, initial-scale=1.0" />
    <link rel="icon" type="image/png" href="../../assets/site/favicon.png" />
    <title>{esc(article["title"])} | 宁波志凡焊材有限公司</title>
    <meta name="description" content="{esc(article["shortAnswer"][:150])}" />
    <meta property="og:title" content="{esc(article["title"])} | 宁波志凡焊材有限公司" />
    <meta property="og:description" content="{esc(article["shortAnswer"][:150])}" />
    <meta name="keywords" content="{esc(article_keywords(article))}" />
    <link rel="stylesheet" href="../../styles.css" />
  </head>
  <body class="article-page qa-article-page">
    <header class="site-header" id="top">
      <nav class="nav" aria-label="文章导航">
        <a class="brand" href="../../index.html#home" aria-label="宁波志凡焊材有限公司首页">
          <img class="brand-logo" src="../../assets/site/zhifan-logo.png" alt="志凡焊材 logo" />
          <span><strong>宁波志凡焊材有限公司</strong><small>ZhiFan Welding Materials</small></span>
        </a>
        <div class="nav-links article-nav-links"><a href="../../index.html#products">产品中心</a><a href="../../index.html#knowledge">焊接操作</a><a class="nav-cta" href="../../index.html#contact">联系我们</a></div>
      </nav>
    </header>
    <main>
      <article class="article-shell qa-shell">
        <a class="back-link" href="../../index.html#knowledge">返回焊接操作</a>
        <div class="article-hero qa-hero">
          <div><p class="eyebrow">{esc(article["category"])} · QA {esc(article["number"])}</p><h1>{esc(article["title"])}</h1><p class="article-generated-time">生成时间：{esc(article["generatedAt"])}</p></div>
          <figure class="article-photo-bg">{article_picture_html()}</figure>
        </div>
        <section class="qa-answer-grid" aria-label="问答内容">
          <article class="qa-answer-card short-answer"><span>短答案</span><h2>先看结论</h2>{paragraphize(article["shortAnswer"])}</article>
          <article class="qa-answer-card long-answer"><span>长答案</span><h2>选型说明</h2>{paragraphize(article["longAnswer"])}</article>
        </section>
        {table_html(article)}
        {note}
        <nav class="article-pager" aria-label="文章翻页"><a href="{esc(prev_link)}">{prev_text}</a><a href="../../index.html#knowledge">返回焊接操作</a><a href="{esc(next_link)}">{next_text}</a><a class="nav-cta" href="../../index.html#contact">联系业务找型号</a></nav>
      </article>
    </main>
    <footer class="footer"><p>© 2026 宁波志凡焊材有限公司</p><span>营业时间：周一至周六 8:00-16:30，周日休息</span><a href="#top">返回顶部</a></footer>
  </body>
</html>
'''
        (out_dir / f'{article["slug"]}.html').write_text(page, encoding="utf-8")


def main():
    doc = Document(str(find_source_docx()))
    articles, categories = extract_articles(doc)
    write_data_module(articles, categories)
    write_pages(articles)
    print(
        json.dumps(
            {
                "articles": len(articles),
                "categories": categories,
                "first": articles[0]["title"],
                "last": articles[-1]["title"],
            },
            ensure_ascii=False,
            indent=2,
        )
    )


if __name__ == "__main__":
    main()
